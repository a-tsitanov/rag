"""Generate the 'Infrastructure sizing' section of the TZ as a .docx.

One-shot script — run once to (re)generate the file.  Output goes to
``docs/specs/2026-04-29-infrastructure-spec.docx``.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm, Pt

OUT = Path(__file__).resolve().parents[1] / "docs" / "specs" / (
    "2026-04-29-infrastructure-spec.docx"
)


def _set_cell(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        _set_cell(t.rows[0].cells[i], h, bold=True)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            _set_cell(t.rows[r_idx].cells[c_idx], val)


def build() -> None:
    doc = Document()

    # ── Title ────────────────────────────────────────────────────────
    title = doc.add_heading(
        "ТЗ. Раздел: Расчёт инфраструктуры", level=0,
    )
    title.runs[0].font.size = Pt(18)
    doc.add_paragraph(
        "Проект: Enterprise Knowledge Base (RAG-платформа на LightRAG). "
        "Документ описывает минимально достаточную и целевую "
        "конфигурацию аппаратных ресурсов для dev/staging/prod.",
    )

    # ── 1. Допущения по нагрузке ─────────────────────────────────────
    doc.add_heading("1. Целевая нагрузка и допущения", level=1)
    _add_table(
        doc,
        headers=["Параметр", "Значение"],
        rows=[
            ["Поток ingestion", "5 000 документов/день (≈ 0.06 док/с)"],
            ["Средний размер документа", "5 МБ (PDF/DOCX/PPTX/TXT)"],
            ["Чанков на документ", "≈ 50 (семантический чанкинг)"],
            ["Поток чанков", "250 000 чанков/день, ~90 М/год"],
            ["RPS поиска", "10–30 QPS пик, 5 QPS средний"],
            ["Latency поиска (hybrid)", "p95 < 500 мс"],
            ["RAGAS faithfulness", "≥ 0.85"],
            ["Uptime API", "99.5%"],
            ["Embedding-модель", "BGE-M3, dim=1024 (через LiteLLM)"],
            ["LLM", "Llama 3.3 70B (Q4/Q8) через LiteLLM"],
            ["Reranker", "BGE-reranker-v2-m3"],
            ["Горизонт хранения", "3 года полных данных, далее архив"],
        ],
    )

    # ── 2. Расчёт хранилища ──────────────────────────────────────────
    doc.add_heading("2. Расчёт хранилища (3 года)", level=1)
    _add_table(
        doc,
        headers=["Компонент", "Формула", "Объём"],
        rows=[
            [
                "Сырые документы (object storage)",
                "5 000 × 5 МБ × 365 × 3",
                "~27 ТБ",
            ],
            [
                "Dense vectors (Milvus)",
                "90M × 1024 × 4 Б × 3",
                "~1.1 ТБ (HNSW +30% overhead → ~1.5 ТБ)",
            ],
            [
                "Sparse vectors / BM25",
                "~10–15% от dense",
                "~150 ГБ",
            ],
            [
                "Knowledge graph (Neo4j)",
                "≈ 50 М узлов, 100 М рёбер + props",
                "~400 ГБ",
            ],
            [
                "PostgreSQL (метаданные, статусы, summary)",
                "~5 КБ/документ × 5.5 М",
                "~30 ГБ",
            ],
            [
                "LightRAG working_dir (KV, doc_status)",
                "~10% от vectors",
                "~150 ГБ",
            ],
            [
                "RabbitMQ + LangFuse + логи",
                "Audit + observability ретеншн 90 д.",
                "~200 ГБ",
            ],
            [
                "ИТОГО (полезный объём)",
                "—",
                "~30 ТБ (рекомендация: 40 ТБ с резервом 30%)",
            ],
        ],
    )

    # ── 3. Сервисы: CPU / RAM / GPU ──────────────────────────────────
    doc.add_heading("3. Ресурсы сервисов (на одну реплику)", level=1)
    _add_table(
        doc,
        headers=["Сервис", "CPU", "RAM", "GPU / VRAM", "Disk (SSD NVMe)"],
        rows=[
            [
                "LLM (Llama 3.3 70B, Q4)",
                "8 vCPU",
                "32 ГБ",
                "1× A100 80 ГБ или 2× L40S 48 ГБ",
                "200 ГБ (модели)",
            ],
            [
                "Embeddings BGE-M3",
                "4 vCPU",
                "16 ГБ",
                "1× T4 16 ГБ или CPU-fallback",
                "20 ГБ",
            ],
            [
                "Reranker BGE-v2-m3",
                "4 vCPU",
                "8 ГБ",
                "Шарит GPU с embeddings",
                "10 ГБ",
            ],
            [
                "LiteLLM proxy",
                "2 vCPU",
                "4 ГБ",
                "—",
                "10 ГБ",
            ],
            [
                "Milvus Distributed (cluster)",
                "16 vCPU суммарно",
                "64 ГБ суммарно",
                "—",
                "1.5 ТБ + резерв",
            ],
            [
                "Neo4j Enterprise",
                "8 vCPU",
                "32 ГБ (heap 16, pagecache 12)",
                "—",
                "500 ГБ",
            ],
            [
                "PostgreSQL 16",
                "4 vCPU",
                "16 ГБ",
                "—",
                "100 ГБ",
            ],
            [
                "RabbitMQ",
                "2 vCPU",
                "4 ГБ",
                "—",
                "50 ГБ",
            ],
            [
                "FastAPI (uvicorn)",
                "4 vCPU × N реплик",
                "8 ГБ × N",
                "—",
                "10 ГБ",
            ],
            [
                "Taskiq worker (ingestion)",
                "8 vCPU × N реплик",
                "16 ГБ × N",
                "—",
                "20 ГБ",
            ],
            [
                "LangFuse (observability)",
                "4 vCPU",
                "8 ГБ",
                "—",
                "100 ГБ",
            ],
        ],
    )

    # ── 4. Минимальная и целевая конфигурации ────────────────────────
    doc.add_heading("4. Сводные конфигурации", level=1)

    doc.add_heading("4.1. Dev / Staging (1 хост)", level=2)
    _add_table(
        doc,
        headers=["Ресурс", "Значение"],
        rows=[
            ["CPU", "32 vCPU (Xeon Gold / EPYC)"],
            ["RAM", "128 ГБ DDR4/DDR5 ECC"],
            ["GPU", "1× A100 40 ГБ или 1× RTX 6000 Ada 48 ГБ"],
            ["Storage", "4 ТБ NVMe SSD (RAID 1)"],
            ["Сеть", "10 Gbps"],
            ["Назначение", "Один хост, все сервисы через docker compose"],
        ],
    )

    doc.add_heading("4.2. Production (кластер)", level=2)
    _add_table(
        doc,
        headers=["Узел", "Кол-во", "Конфигурация"],
        rows=[
            [
                "GPU-нода (LLM)",
                "2 (active + standby)",
                "32 vCPU, 256 ГБ RAM, 1× A100 80 ГБ, 2 ТБ NVMe",
            ],
            [
                "GPU-нода (embed/rerank)",
                "2",
                "16 vCPU, 64 ГБ RAM, 1× L4 24 ГБ, 500 ГБ NVMe",
            ],
            [
                "Milvus / Neo4j storage-нода",
                "3",
                "32 vCPU, 128 ГБ RAM, 8 ТБ NVMe (RAID 10)",
            ],
            [
                "App-нода (API + worker)",
                "3",
                "16 vCPU, 64 ГБ RAM, 500 ГБ SSD",
            ],
            [
                "Infra-нода (Postgres + Rabbit + LangFuse)",
                "2 (HA-пара)",
                "16 vCPU, 64 ГБ RAM, 2 ТБ NVMe",
            ],
            [
                "Object storage (raw docs)",
                "S3-совместимое (MinIO HA)",
                "40+ ТБ полезного объёма, репликация ×3",
            ],
            [
                "Сеть",
                "—",
                "25 Gbps между нодами, 10 Gbps uplink",
            ],
        ],
    )
    doc.add_paragraph(
        "ИТОГО prod-кластер: 12 узлов, ~256 vCPU, ~1 ТБ RAM, 4× GPU, "
        "~50 ТБ NVMe.",
    )

    # ── 5. Сеть и безопасность ───────────────────────────────────────
    doc.add_heading("5. Сетевые и эксплуатационные требования", level=1)
    doc.add_paragraph(
        "• Внутренняя сеть L2/L3 ≥ 10 Gbps между Milvus / Neo4j / GPU-нодами; "
        "TLS 1.3 на всех публичных эндпоинтах.\n"
        "• Доступ к API только через reverse-proxy (nginx / Traefik) с rate-limit "
        "и аутентификацией по X-API-Key.\n"
        "• LiteLLM-прокси изолирован во внутренней сети, наружу не публикуется.\n"
        "• Бэкапы: PG + Neo4j ежедневно, Milvus snapshot — еженедельно, "
        "хранение 30 дней.\n"
        "• Мониторинг: Prometheus + Grafana + LangFuse, алерты в Slack/PagerDuty.\n"
        "• SLO: API uptime 99.5%, ingestion success rate ≥ 99%, "
        "search p95 < 500 мс.",
    )

    # ── 6. Масштабирование ───────────────────────────────────────────
    doc.add_heading("6. План масштабирования", level=1)
    doc.add_paragraph(
        "• До 10 000 док/день: горизонтальный +1 worker-нода и +1 Milvus QueryNode.\n"
        "• До 50 000 док/день: шардирование Milvus по департаменту, выделенный "
        "Neo4j Causal Cluster (3+ core), увеличение GPU-пула до 4×A100.\n"
        "• Embedding throughput скейлится репликами BGE-M3 за LiteLLM "
        "load-balancer'ом.",
    )

    # Document-wide font tweak
    for p in doc.paragraphs:
        for run in p.runs:
            if not run.font.size:
                run.font.size = Pt(11)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
